#!/usr/bin/env python3
"""Generate the HubSpot Career Chart Implementation Guide as .docx"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size, color in [('Heading 1', 22, (0x26, 0x1f, 0x46)), ('Heading 2', 16, (0x26, 0x1f, 0x46)), ('Heading 3', 13, (0xf0, 0x92, 0x21))]:
    h = doc.styles[level]
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(*color)
    h.font.bold = True
    h.paragraph_format.space_before = Pt(18 if level != 'Heading 1' else 24)
    h.paragraph_format.space_after = Pt(8)

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('YOGAPROS')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)
run.font.bold = True
run.font.name = 'Calibri'

doc.add_paragraph('')

title2 = doc.add_paragraph()
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title2.add_run('Career Chart')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)
run.font.bold = True
run.font.name = 'Calibri'

title3 = doc.add_paragraph()
title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title3.add_run('HubSpot Implementation Guide')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0xf0, 0x92, 0x21)
run.font.bold = False
run.font.name = 'Calibri'

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Step-by-step guide to implementing the YogaPros Career Chart\nin HubSpot CMS with full data capture and workflow automation')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'Calibri'

doc.add_paragraph('')
doc.add_paragraph('')

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'March 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Contents', level=1)

toc_items = [
    ('1.', 'Overview', 'What the Career Chart is and how it works'),
    ('2.', 'Architecture', 'How the pieces fit together in HubSpot'),
    ('3.', 'Step 1: Create Custom Contact Properties', '14 properties to create (point and click)'),
    ('4.', 'Step 2: Create the Hidden HubSpot Form', 'Drag-and-drop form builder'),
    ('5.', 'Step 3: Add the Career Chart Page', 'Paste the code into HubSpot CMS'),
    ('6.', 'Step 4: Configure Your Portal ID & Form GUID', 'Two values to plug in'),
    ('7.', 'Step 5: Build Workflows', 'Automated nurture sequences based on scores'),
    ('8.', 'Step 6: Test Everything', 'Testing checklist'),
    ('', '', ''),
    ('A.', 'Appendix: Custom Properties Reference', 'Full list of all 14 properties'),
    ('B.', 'Appendix: Complete Quiz Code', 'The full HTML/CSS/JS to paste'),
]

for num, title_text, desc in toc_items:
    if not num:
        continue
    p = doc.add_paragraph()
    run = p.add_run(f'{num} {title_text}')
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)
    if desc:
        run2 = p.add_run(f'  \u2014  {desc}')
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ============================================================
# 1. OVERVIEW
# ============================================================
doc.add_heading('1. Overview', level=1)

doc.add_paragraph(
    'The Career Chart is a 12-question interactive quiz that maps where a yoga teacher '
    'sits across three pillars: Teaching, Business, and Wellbeing & Mindset (plus Safety & Protection). '
    'It collects their name, email, and role before starting, then captures every answer and calculated score '
    'into HubSpot as contact properties.'
)

doc.add_heading('What it does', level=2)

items = [
    'Captures name, email, and role (Teacher / Trainer / Aspiring)',
    'Asks 12 scored questions across 4 categories',
    'Calculates pillar scores, percentages, and tier levels instantly in-browser',
    'Shows personalised results page with pillar-by-pillar breakdown',
    'Silently submits ALL data to HubSpot in the background',
    'Creates/updates a HubSpot contact with 14 custom properties',
    'Triggers automated workflows based on scores and answers',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_heading('What Aimee does vs what the code does', level=2)

table = doc.add_table(rows=4, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'

headers = ['Aimee (point and click)', 'Code (already built)']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

tasks = [
    ('Create 14 custom contact properties', 'Calculates scores from quiz answers'),
    ('Create a hidden HubSpot form', 'Submits data to HubSpot via Forms API'),
    ('Paste code into HubSpot CMS page\nBuild follow-up workflows\nTest', 'Renders the full quiz UI and results page'),
]
for i, (aimee, code) in enumerate(tasks):
    table.rows[i+1].cells[0].text = aimee
    table.rows[i+1].cells[1].text = code

doc.add_page_break()

# ============================================================
# 2. ARCHITECTURE
# ============================================================
doc.add_heading('2. Architecture', level=1)

doc.add_paragraph('Here is how the data flows:')

flow_steps = [
    'Teacher lands on Career Chart page',
    'Enters name, email, and role \u2192 clicks "Start"',
    'Answers 12 questions (all happens in their browser, no server calls)',
    'JavaScript instantly calculates scores and shows personalised results',
    'In the background, a single API call sends everything to HubSpot',
    'HubSpot creates or updates the contact with all 14 custom properties',
    'Workflows trigger based on scores (nurture sequences, insurance alerts, etc.)',
]
for i, step in enumerate(flow_steps):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i+1}: ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xf0, 0x92, 0x21)
    p.add_run(step)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Important: ')
run.font.bold = True
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
p.add_run('The teacher sees their results instantly. They never wait for HubSpot. The API call happens invisibly after results are displayed.')

doc.add_page_break()

# ============================================================
# 3. CREATE CUSTOM PROPERTIES
# ============================================================
doc.add_heading('3. Step 1: Create Custom Contact Properties', level=1)

doc.add_paragraph('You need to create 14 custom properties on the Contact object in HubSpot. These store the quiz data on each contact record.')

doc.add_heading('How to get there', level=2)

nav_steps = [
    'Go to Settings (gear icon, top right)',
    'Click Properties in the left sidebar (under Data Management)',
    'Make sure "Contact properties" is selected at the top',
    'Click "Create property" (orange button, top right)',
]
for step in nav_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Create a property group first', level=2)

doc.add_paragraph('Before creating individual properties, create a group to keep them organised:')

group_steps = [
    'In the "Create property" screen, under Group, click "Create a new group"',
    'Name it: Career Chart',
    'This keeps all quiz data together and easy to find',
]
for step in group_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Properties to create', level=2)

doc.add_paragraph('Create each of the following properties. For each one:')

create_steps = [
    'Object type: Contact',
    'Group: Career Chart (the group you just created)',
    'Label: as shown below',
    'Internal name: will auto-generate from label (shown in brackets)',
    'Field type: as shown below',
]
for step in create_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph('')

# Properties table
props = [
    ('1', 'Career Chart Score', 'career_chart_score', 'Number', 'Overall score (0-100)'),
    ('2', 'Career Chart Tier', 'career_chart_tier', 'Single-line text', 'Foundation / Building / Professional'),
    ('3', 'Teaching Score', 'cc_teaching_score', 'Number', 'Teaching pillar % (0-100)'),
    ('4', 'Teaching Level', 'cc_teaching_level', 'Single-line text', 'e.g. "Experienced Teacher"'),
    ('5', 'Business Score', 'cc_business_score', 'Number', 'Business pillar % (0-100)'),
    ('6', 'Business Level', 'cc_business_level', 'Single-line text', 'e.g. "Building a Brand"'),
    ('7', 'Mindset Score', 'cc_mindset_score', 'Number', 'Mindset pillar % (0-100)'),
    ('8', 'Mindset Level', 'cc_mindset_level', 'Single-line text', 'e.g. "Awareness"'),
    ('9', 'Safety Score', 'cc_safety_score', 'Number', 'Safety pillar % (0-100)'),
    ('10', 'Safety Level', 'cc_safety_level', 'Single-line text', 'e.g. "Fully Protected"'),
    ('11', 'CC Q: Has Insurance', 'cc_has_insurance', 'Single-line text', 'Answer to insurance question'),
    ('12', 'CC Q: Years Teaching', 'cc_years_teaching', 'Single-line text', 'Answer to experience question'),
    ('13', 'CC Q: Income Model', 'cc_income_model', 'Single-line text', 'Answer to income question'),
    ('14', 'CC: Date Completed', 'cc_date_completed', 'Date picker', 'When they took the quiz'),
]

table = doc.add_table(rows=len(props)+1, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['#', 'Label', 'Internal Name', 'Field Type', 'Description']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(9)

for row_idx, prop in enumerate(props):
    for col_idx, val in enumerate(prop):
        cell = table.rows[row_idx+1].cells[col_idx]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Tip: ')
run.font.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
p.add_run('The internal names shown above are suggestions. HubSpot will auto-generate internal names from the labels. '
          'Just note down the exact internal names HubSpot creates \u2014 you will need them in Step 4.')

doc.add_page_break()

# ============================================================
# 4. CREATE HIDDEN FORM
# ============================================================
doc.add_heading('4. Step 2: Create the Hidden HubSpot Form', level=1)

doc.add_paragraph('This form is never shown to users. The quiz JavaScript submits data to it via API. But HubSpot needs the form to exist in order to accept submissions.')

doc.add_heading('How to create it', level=2)

form_steps = [
    ('Go to Marketing \u2192 Forms', ''),
    ('Click "Create form" \u2192 choose "Embedded form"', ''),
    ('Name it: "Career Chart Submission" (internal only, users never see this)', ''),
    ('Add ALL of the following fields by dragging them in:', 'First name (built-in)\nEmail (built-in)\nAll 14 custom properties you created in Step 1'),
    ('Turn OFF the form\'s built-in submit button', 'Under Options \u2192 toggle off "Submit button"'),
    ('Under Options, turn OFF all "follow-up" options', 'No thank you page, no email notification \u2014 the quiz handles all of this'),
    ('Click "Publish"', ''),
    ('After publishing, click "Share" or "Embed"', ''),
    ('Copy the Form GUID from the embed code', 'It looks like: a1b2c3d4-e5f6-7890-abcd-ef1234567890'),
    ('Also note your Portal ID', 'Visible in the embed code or in your HubSpot URL (the number after "app/")'),
]

for i, (step, detail) in enumerate(form_steps):
    p = doc.add_paragraph()
    run = p.add_run(f'{i+1}. ')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xf0, 0x92, 0x21)
    run2 = p.add_run(step)
    run2.font.bold = True
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run3 = p2.add_run(detail)
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('You now have two values you need for Step 4:')
run.font.bold = True

vals = [
    'Portal ID \u2014 e.g. 12345678',
    'Form GUID \u2014 e.g. a1b2c3d4-e5f6-7890-abcd-ef1234567890',
]
for v in vals:
    doc.add_paragraph(v, style='List Bullet')

doc.add_page_break()

# ============================================================
# 5. ADD THE PAGE
# ============================================================
doc.add_heading('5. Step 3: Add the Career Chart Page', level=1)

doc.add_paragraph('There are two ways to add the quiz to HubSpot. Choose the one that suits your setup:')

doc.add_heading('Option A: Custom Module (Recommended)', level=2)

option_a = [
    'Go to Marketing \u2192 Website \u2192 Website Pages',
    'Create a new page or open an existing one',
    'In the page editor, add a "Custom HTML" or "Rich text" module',
    'Switch to "Source code" / HTML view',
    'Paste the complete quiz code from Appendix B',
    'Save and preview',
]
for i, step in enumerate(option_a):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_heading('Option B: HubSpot Custom Module (Advanced)', level=2)

doc.add_paragraph('If you want it as a reusable drag-and-drop module:')

option_b = [
    'Go to Marketing \u2192 Files and Templates \u2192 Design Tools',
    'Create a new Custom Module',
    'Paste the HTML into the HTML section',
    'Paste the CSS into the CSS section',
    'Paste the JavaScript into the JS section',
    'Save and the module becomes available in any page editor',
]
for i, step in enumerate(option_b):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Note: ')
run.font.bold = True
run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)
p.add_run('Option A is simpler and faster. Option B is better if you want to reuse the quiz on multiple pages or if your HubSpot theme requires it.')

doc.add_page_break()

# ============================================================
# 6. CONFIGURE PORTAL ID & FORM GUID
# ============================================================
doc.add_heading('6. Step 4: Configure Your Portal ID & Form GUID', level=1)

doc.add_paragraph('In the quiz code (Appendix B), find these two lines near the top of the JavaScript section:')

code_para = doc.add_paragraph()
code_para.paragraph_format.left_indent = Cm(1)
run = code_para.add_run("const HUBSPOT_PORTAL_ID = 'YOUR_PORTAL_ID_HERE';\nconst HUBSPOT_FORM_GUID = 'YOUR_FORM_GUID_HERE';")
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

doc.add_paragraph('')
doc.add_paragraph('Replace these with the values you noted in Step 2:')

code_para2 = doc.add_paragraph()
code_para2.paragraph_format.left_indent = Cm(1)
run = code_para2.add_run("const HUBSPOT_PORTAL_ID = '12345678';  // your portal ID\nconst HUBSPOT_FORM_GUID = 'a1b2c3d4-e5f6-7890-abcd-ef1234567890';  // your form GUID")
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_heading('How to find your Portal ID', level=2)

portal_steps = [
    'Log into HubSpot',
    'Look at the URL in your browser \u2014 it contains your portal ID',
    'Example: app.hubspot.com/contacts/12345678 \u2192 Portal ID is 12345678',
    'Or: Settings \u2192 Account Defaults \u2192 Account Information',
]
for step in portal_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('How to find your Form GUID', level=2)

guid_steps = [
    'Go to Marketing \u2192 Forms',
    'Click on "Career Chart Submission" (the form you created)',
    'Click "Share" or look at the embed code',
    'The GUID is the long alphanumeric string in the code',
    'Or look at the URL when editing the form \u2014 it\'s the ID at the end',
]
for step in guid_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('Property name mapping', level=2)

doc.add_paragraph('The code maps quiz data to HubSpot property internal names. If HubSpot generated different internal names than the ones below, you need to update the code to match.')

doc.add_paragraph('Find the PROPERTY_MAP object in the JavaScript and update the values to match your actual HubSpot internal property names:')

code_para3 = doc.add_paragraph()
code_para3.paragraph_format.left_indent = Cm(1)
run = code_para3.add_run(
    "const PROPERTY_MAP = {\n"
    "    score: 'career_chart_score',\n"
    "    tier: 'career_chart_tier',\n"
    "    teaching_score: 'cc_teaching_score',\n"
    "    teaching_level: 'cc_teaching_level',\n"
    "    business_score: 'cc_business_score',\n"
    "    business_level: 'cc_business_level',\n"
    "    mindset_score: 'cc_mindset_score',\n"
    "    mindset_level: 'cc_mindset_level',\n"
    "    safety_score: 'cc_safety_score',\n"
    "    safety_level: 'cc_safety_level',\n"
    "    has_insurance: 'cc_has_insurance',\n"
    "    years_teaching: 'cc_years_teaching',\n"
    "    income_model: 'cc_income_model',\n"
    "    date_completed: 'cc_date_completed'\n"
    "};"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)

doc.add_page_break()

# ============================================================
# 7. WORKFLOWS
# ============================================================
doc.add_heading('7. Step 5: Build Workflows', level=1)

doc.add_paragraph('This is where the data becomes powerful. Each quiz taker gets scored and categorised, so you can send them exactly the right follow-up.')

doc.add_heading('Recommended Workflows', level=2)

# Workflow 1
doc.add_heading('Workflow 1: Score-Based Nurture Sequences', level=3)

p = doc.add_paragraph()
run = p.add_run('Trigger: ')
run.font.bold = True
p.add_run('Contact property "Career Chart Score" is known')

doc.add_paragraph('')

tiers = [
    ('Score < 40%', 'Foundation Sequence', 'Mindset-first content. These teachers are stuck in scarcity. Send content about the abundance shift, identity as a professional, "income sustains my teaching" messaging. Do NOT lead with features or pricing.'),
    ('Score 40\u201374%', 'Building Sequence', 'Business + visibility focus. These teachers know they need to professionalise but need help with the "how." Lead with Live CV, discoverability, pricing confidence. Show them what the professional ecosystem looks like.'),
    ('Score 75%+', 'Professional Sequence', 'Direct Professional membership invite. These teachers are ready. Lead with identity confirmation ("You already operate at this level"), community, and the tangible tools. Short sequence, clear CTA.'),
]

for score_range, seq_name, desc in tiers:
    p = doc.add_paragraph()
    run = p.add_run(f'{score_range} \u2192 "{seq_name}"')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# Workflow 2
doc.add_heading('Workflow 2: No Insurance Alert', level=3)

p = doc.add_paragraph()
run = p.add_run('Trigger: ')
run.font.bold = True
p.add_run('"CC Q: Has Insurance" contains "No" or "didn\'t know"')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Action: ')
run.font.bold = True
p.add_run('Send immediate insurance-focused email. Lead with protection, not fear. "You deserve to teach with confidence. Professional membership includes the most comprehensive insurance in the industry \u2014 it\'s one of the first things members tell us they value."')

# Workflow 3
doc.add_heading('Workflow 3: Mindset Blocker', level=3)

p = doc.add_paragraph()
run = p.add_run('Trigger: ')
run.font.bold = True
p.add_run('"Mindset Score" is less than 40')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Action: ')
run.font.bold = True
p.add_run('Enrol in scarcity-to-abundance content series. 5-7 emails over 2 weeks. Frame as identity content, not sales. "The shift from \'I shouldn\'t charge\' to \'my expertise has value\' is the single biggest unlock in a yoga teacher\'s career."')

# Workflow 4
doc.add_heading('Workflow 4: Aspiring Teachers', level=3)

p = doc.add_paragraph()
run = p.add_run('Trigger: ')
run.font.bold = True
p.add_run('Role = "Aspiring Yoga Teacher"')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Action: ')
run.font.bold = True
p.add_run('Different sequence entirely. These are pre-qualification. Guide them toward YogaPros-accredited training schools. "Start right. Choose a training school that meets the highest standards from day one."')

doc.add_heading('How to create a workflow', level=2)

wf_steps = [
    'Go to Automation \u2192 Workflows',
    'Click "Create workflow" \u2192 "From scratch"',
    'Choose "Contact-based" workflow',
    'Set the enrollment trigger (e.g., "Career Chart Score is known")',
    'Add an If/then branch based on the score value',
    'Add email actions for each branch',
    'Turn on the workflow',
]
for i, step in enumerate(wf_steps):
    doc.add_paragraph(f'{i+1}. {step}')

doc.add_page_break()

# ============================================================
# 8. TESTING
# ============================================================
doc.add_heading('8. Step 6: Test Everything', level=1)

doc.add_paragraph('Before going live, test the full flow:')

tests = [
    ('Open the Career Chart page', 'Check it loads correctly, looks right on desktop and mobile'),
    ('Fill in name, email, role', 'Use a test email you can check'),
    ('Complete all 12 questions', 'Check progress bar works, back button works'),
    ('View results page', 'Check scores display correctly, personalised content shows'),
    ('Check HubSpot contact', 'Go to Contacts, find your test email \u2014 all 14 properties should be populated'),
    ('Verify workflow enrollment', 'Check the test contact was enrolled in the correct workflow based on score'),
    ('Test on mobile', 'Open the page on your phone and complete the quiz'),
    ('Test with different scores', 'Take it 3 times with low/medium/high answers to verify all workflow branches'),
]

for i, (test, detail) in enumerate(tests):
    p = doc.add_paragraph()
    run = p.add_run(f'\u2610  {test}')
    run.font.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run(detail)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Troubleshooting: ')
run.font.bold = True
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
p.add_run('If data is not appearing in HubSpot, open your browser\'s Developer Tools (F12 or Cmd+Option+I), go to the Console tab, and look for error messages when you submit the quiz. The most common issue is mismatched property names between the code and HubSpot.')

doc.add_page_break()

# ============================================================
# APPENDIX A: PROPERTIES REFERENCE
# ============================================================
doc.add_heading('Appendix A: Custom Properties Reference', level=1)

doc.add_paragraph('Full specification for each property:')

for prop in props:
    num, label, internal, field_type, desc = prop
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {label}')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x26, 0x1f, 0x46)

    details = [
        f'Internal name: {internal}',
        f'Field type: {field_type}',
        f'Group: Career Chart',
        f'Description: {desc}',
    ]
    for d in details:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1)
        run2 = p2.add_run(d)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ============================================================
# APPENDIX B: COMPLETE CODE
# ============================================================
doc.add_heading('Appendix B: Complete Career Chart Code', level=1)

doc.add_paragraph('Copy and paste the entire code block below into your HubSpot page. The code is self-contained \u2014 all HTML, CSS, and JavaScript in a single block.')

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('IMPORTANT: ')
run.font.bold = True
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
p.add_run('Before pasting, update the two configuration values at the top of the <script> section (marked with YOUR_PORTAL_ID_HERE and YOUR_FORM_GUID_HERE) with your actual HubSpot Portal ID and Form GUID from Step 4.')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('The complete code is provided as a separate HTML file: ')
run.font.size = Pt(11)
run2 = p.add_run('career-chart-hubspot.html')
run2.font.bold = True
run2.font.size = Pt(11)
run2.font.color.rgb = RGBColor(0xf0, 0x92, 0x21)

doc.add_paragraph('')
doc.add_paragraph('This file should be attached alongside this guide. Open it in a text editor (e.g. TextEdit on Mac \u2014 make sure to use Format \u2192 Make Plain Text first, or use a code editor like VS Code), copy all contents, and paste into HubSpot.')

doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run('Note on the code file: ')
run.font.bold = True
p.add_run('The HTML file can also be opened in a web browser to preview the quiz locally (the HubSpot submission will fail without valid credentials, but the quiz UI and results will work). This is useful for reviewing the look and feel before deploying to HubSpot.')

# ============================================================
# SAVE
# ============================================================
output_path = '/Users/bcm/Library/Mobile Documents/com~apple~CloudDocs/Claude Code/root-and-rise/YogaPros_Career_Chart_HubSpot_Guide.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
